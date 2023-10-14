
import pytest
import os
import time
from dragonfly import Window
from pynput.mouse import Button, Controller
from pynput.keyboard import Controller as KController
from docx import Document
from docx.shared import Inches
import pyautogui

# Функция генерирует скриншот с docx word файлом в котором описаны этапы баг-репорта согласно тех. заданию
def write_bugreport(testCase='', caseStep='', expectResult='', actualResult=''):

    document = Document()

    pyautogui.screenshot(f'bug-{testCase}-{caseStep}.png')

    table = document.add_table(rows=5, cols=2, style='TableGrid')

    table.cell(0, 0).add_paragraph(text='Тест-кейс:')
    table.cell(1, 0).add_paragraph(text='Шаг тест-кейса:')
    table.cell(2, 0).add_paragraph(text='Ожидаемый результат:')
    table.cell(3, 0).add_paragraph(text='Фактический результат:')
    table.cell(4, 0).add_paragraph(text='Окружение:')

    table.cell(0, 1).add_paragraph().add_run(testCase).bold = True
    table.cell(1, 1).add_paragraph().add_run(caseStep).bold = True
    table.cell(2, 1).add_paragraph().add_run(expectResult).bold = True
    table.cell(3, 1).add_paragraph().add_run(actualResult).bold = True
    table.cell(4, 1).add_paragraph().add_run('Windows 10 Pro - 1909 (18363.1556); Яндекс Браузер - 23.9.2.888 (64-bit)').bold = True

    document.add_picture(f'bug-{testCase}-{caseStep}.png', width=Inches(7), height=Inches(5))
    
    return document.save(F'bugreport-{testCase}-{caseStep}.docx')

# Шаг 1: запустить браузер
def test_step1():
    """Шаг 1: Ожидаемый результат - браузер запускается"""

    os.startfile(F'{os.getenv("LOCALAPPDATA")}\\Yandex\\YandexBrowser\\Application\\browser.exe')

    time.sleep(10)

    # Сравниваем количетсво окон зависимых от browser.exe (Yandex Browser); 0 - окна не появились, браузер не запустился
    try:
        assert len(Window.get_matching_windows(executable='browser.exe')) != 0

    except AssertionError:

        write_bugreport('2','1','Браузер запускается','Браузер не запустился!')

        assert False, 'Шаг 1 - ПРОВАЛ, Причина: браузер не запустился'

# Шаг 2: нажать на кнопку "+" в боковой панели
def test_step2():
    """Шаг 2: Ожидаемый результат - появляется баллун"""

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(1)

    mouse.move(24, 933)

    time.sleep(1)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(3)

# Шаг 3: нажать на кнопку "Добавить сайт как приложение"
def test_step3():
    """Шаг 3: Ожидаемый результат - появляется меню добавления приложения"""

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(1)

    mouse.move(236, 993)

    time.sleep(1)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(3)

# Шаг 4: ввести в строку pikabu.ru
def test_step4():
    """Шаг 4: Ожидаемый результат - появляются подсказки"""

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(1)

    mouse.move(221, 165)

    time.sleep(1)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(1)

    keyboard = KController()

    string_to_write = 'pikabu.ru'

    for char in string_to_write:
        keyboard.press(char)
        keyboard.release(char)
        time.sleep(0.15)

    time.sleep(3)

# Шаг 5: Выбрать первый вариант и кликнуть на него мышью
def test_step5():
    """Шаг 5: Ожидаемый результат - приложение устанавливается: в боковой панели, появляется его иконка"""

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(1)

    mouse.move(207, 207)

    time.sleep(1)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(20)

# Шаг 6: Кликнуть на иконку приложения в боковой панели
def test_step6():
    """Шаг 6: Ожидаемый результат - приложение запускается в отдельном окне"""

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(2)

    mouse.move(25, 870)

    time.sleep(2)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(5)

    # перебираем словарь и конвертируем id окон в integer для параметра Window и сравниваем кол-во окон
    try:
        window_title = ''

        for i in Window.get_matching_windows(executable='browser.exe'):

            convertstring = str.replace(str(i), 'Win32Window(handle=', '')

            convertstring = int(str.replace(convertstring, ')', ''))

            if 'Пикабу' in Window(convertstring).title:

                window_title = Window(convertstring).title

        assert len(Window.get_matching_windows(executable='browser.exe', title=window_title)) != 0

    except AssertionError:

        write_bugreport('2','6','приложение запускается в отдельном окне','окно не запустилось')

        assert False, 'Шаг 6 - ПРОВАЛ, Причина: окно не запустилось'

# Шаг 7: прокликать по всем ссылкам в шапке сайта
def test_step7():
    """Шаг 7: Ожидаемый результат - открытие новых страниц происходит внутри окна приложения, а не внутри основного окна браузера"""

    original_browser_name = ''

    # открываем отдельное окно с Пикабу в полный экран т.к. позиция окна может быть случайна что делает невозможным симуляцию мышки по пиксель-координатам
    for i in Window.get_matching_windows(executable='browser.exe'):

        convertstring = str.replace(str(i), 'Win32Window(handle=', '')

        convertstring = int(str.replace(convertstring, ')', ''))

        if 'Пикабу' in Window(convertstring).title:

            Window(convertstring).maximize()
        
        if ' — Яндекс' in Window(convertstring).title:

            original_browser_name = Window(convertstring).title

    time.sleep(2)

    mouse = Controller()

    mouse.position = (541,71)

    # Перебираем увеличение координат мыши на 80 пикселей 6 раз
    for step in range(6):

        mouse.move(80, 0)

        time.sleep(2)

        mouse.press(Button.left)
        time.sleep(0.36)
        mouse.release(Button.left)

        time.sleep(5)

        # Сравниваем первое название окна browser.exe, если при клике влияние оказывается на основное окно web-браузера, значит и название окна будет меняться
        try:
            for i in Window.get_matching_windows(executable='browser.exe'):

                convertstring = str.replace(str(i), 'Win32Window(handle=', '')

                convertstring = int(str.replace(convertstring, ')', ''))
            
                if ' — Яндекс' in Window(convertstring).title:

                    assert original_browser_name == Window(convertstring).title

        except AssertionError:

            write_bugreport('2','7','открытие новых страниц происходит внутри окна приложения, а не внутри основного окна браузера',
                            'страницы открываются внутри основного окна браузера')

            assert False, 'Шаг 7 - ПРОВАЛ, Причина: страницы открываются внутри основного окна браузера'

    # Закрываем окна для подготовки к тест-кейсу № 2
    for i in Window.get_matching_windows(executable='browser.exe'):

        convertstring = str.replace(str(i), 'Win32Window(handle=', '')

        convertstring = int(str.replace(convertstring, ')', ''))

        Window(convertstring).close()

    time.sleep(10)