
import pytest
import os
import time
from dragonfly import Window
from pynput.mouse import Button, Controller
from docx import Document
from docx.shared import Inches
import pyautogui

# Функция генерирует скриншот с docx word файлом в котором описаны этапы баг-репорта согласно тех. заданию
def write_bugreport(testCase='', caseStep='', expectResult='', actualResult=''):

    document = Document()

    pyautogui.screenshot(f'bug-{testCase}-{caseStep}.png')

    table = document.add_table(rows=5, cols=2, style='TableGrid')

    table.cell(0, 0).add_paragraph(text='Тест-кейс:')
    table.cell(1, 0).add_paragraph(text='Шаг тест-кейса:')
    table.cell(2, 0).add_paragraph(text='Ожидаемый результат:')
    table.cell(3, 0).add_paragraph(text='Фактический результат:')
    table.cell(4, 0).add_paragraph(text='Окружение:')

    table.cell(0, 1).add_paragraph().add_run(testCase).bold = True
    table.cell(1, 1).add_paragraph().add_run(caseStep).bold = True
    table.cell(2, 1).add_paragraph().add_run(expectResult).bold = True
    table.cell(3, 1).add_paragraph().add_run(actualResult).bold = True
    table.cell(4, 1).add_paragraph().add_run('Windows 10 Pro - 1909 (18363.1556); Яндекс Браузер - 23.9.2.888 (64-bit)').bold = True

    document.add_picture(f'bug-{testCase}-{caseStep}.png', width=Inches(7), height=Inches(5))
    
    return document.save(F'bugreport-{testCase}-{caseStep}.docx')

# Шаг 1: запустить браузер
def test_step1():
    """Шаг 1: Ожидаемый результат - браузер запускается"""

    os.startfile(F'{os.getenv("LOCALAPPDATA")}\\Yandex\\YandexBrowser\\Application\\browser.exe')

    time.sleep(10)

    # Сравниваем количетсво окон зависимых от browser.exe (Yandex Browser); 0 - окна не появились, браузер не запустился
    try:
        assert len(Window.get_matching_windows(executable='browser.exe')) != 0

    except AssertionError:

        write_bugreport('1','1','Браузер запускается','Браузер не запустился!')

        assert False, 'Шаг 1 - ПРОВАЛ, Причина: браузер не запустился'

# Шаг 2: запустить приложение Почта из боковой панели
def test_step2():
    """Шаг 2: Ожидаемый результат - приложение запускается в отдельном окне"""

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(2)

    mouse.move(25, 902)

    time.sleep(2)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(5)

    # Сравниваем количетсво окон зависимых от browser.exe (Yandex Browser) с названием "Яндекс Почта"; 0 - окна не появились
    try:
        assert len(Window.get_matching_windows(executable='browser.exe', title='Яндекс Почта')) != 0

    except AssertionError:

        write_bugreport('1','2','приложение запускается в отдельном окне','окно не запустилось')

        assert False, 'Шаг 2 - ПРОВАЛ, Причина: окно не запустилось'

    # Закрываем окна для подготовки к тест-кейсу № 2
    for i in Window.get_matching_windows(executable='browser.exe'):

        convertstring = str.replace(str(i), 'Win32Window(handle=', '')

        convertstring = int(str.replace(convertstring, ')', ''))

        Window(convertstring).close()

    time.sleep(10)
