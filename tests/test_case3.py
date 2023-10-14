
import pytest
import os
import time
from dragonfly import Window
from pynput.mouse import Button, Controller
from pynput.keyboard import Controller as KController
from pynput.keyboard import Key
from docx import Document
from docx.shared import Inches
import pyautogui

# Функция генерирует скриншот с docx word файлом в котором описаны этапы баг-репорта согласно тех. заданию
def write_bugreport(testCase='', caseStep='', expectResult='', actualResult=''):

    document = Document()

    pyautogui.screenshot(f'bug-{testCase}-{caseStep}.png')

    table = document.add_table(rows=5, cols=2, style='TableGrid')

    table.cell(0, 0).add_paragraph(text='Тест-кейс:')
    table.cell(1, 0).add_paragraph(text='Шаг тест-кейса:')
    table.cell(2, 0).add_paragraph(text='Ожидаемый результат:')
    table.cell(3, 0).add_paragraph(text='Фактический результат:')
    table.cell(4, 0).add_paragraph(text='Окружение:')

    table.cell(0, 1).add_paragraph().add_run(testCase).bold = True
    table.cell(1, 1).add_paragraph().add_run(caseStep).bold = True
    table.cell(2, 1).add_paragraph().add_run(expectResult).bold = True
    table.cell(3, 1).add_paragraph().add_run(actualResult).bold = True
    table.cell(4, 1).add_paragraph().add_run('Windows 10 Pro - 1909 (18363.1556); Яндекс Браузер - 23.9.2.888 (64-bit)').bold = True

    document.add_picture(f'bug-{testCase}-{caseStep}.png', width=Inches(7), height=Inches(5))
    
    return document.save(F'bugreport-{testCase}-{caseStep}.docx')

# Шаг 1: запустить браузер
def test_step1():
    """Шаг 1: Ожидаемый результат - браузер запускается"""

    os.startfile(F'{os.getenv("LOCALAPPDATA")}\\Yandex\\YandexBrowser\\Application\\browser.exe')

    time.sleep(10)

    try:
        assert len(Window.get_matching_windows(executable='browser.exe')) != 0

    except AssertionError:

        write_bugreport('3','1','Браузер запускается','Браузер не запустился!')

        assert False, 'Шаг 1 - ПРОВАЛ, Причина: браузер не запустился'

# Шаг 2: в настройках браузера, в разделе Интерфейс, выбрать тему под названием "Прозрачная"
def test_step2():
    """Шаг 2: Ожидаемый результат - тема применилась"""

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(2)

    mouse.move(1765, 19)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(2)

    mouse.position = (0,0)

    time.sleep(5)

    mouse.move(1543, 220)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(2)

    mouse.position = (0,0)

    time.sleep(10)

    mouse.move(426, 267)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(2)

    mouse.position = (0,0)

    time.sleep(2)

    mouse.move(801, 553)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(10)

# Шаг 3: перейти на сайт https://vk.com/
def test_step3():
    """Шаг 3: Ожидаемый результат - сайт загружается"""

    original_browser_name = ''

    # Запоминаем название вкладки до перехода на сайт
    for i in Window.get_matching_windows(executable='browser.exe'):

        convertstring = str.replace(str(i), 'Win32Window(handle=', '')

        convertstring = int(str.replace(convertstring, ')', ''))
        
        if ' — Яндекс' in Window(convertstring).title:

            original_browser_name = Window(convertstring).title

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(2)

    mouse.move(248, 59)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    keyboard = KController()

    time.sleep(2)

    string_to_write = 'https://vk.com/'

    for char in string_to_write:
        keyboard.press(char)
        keyboard.release(char)
        time.sleep(0.15)

    keyboard.press(Key.enter)
    keyboard.release(Key.enter)

    time.sleep(20)

    # Сравниваем название вкладок до перехода на сайт и после
    try:
        for i in Window.get_matching_windows(executable='browser.exe'):

            convertstring = str.replace(str(i), 'Win32Window(handle=', '')

            convertstring = int(str.replace(convertstring, ')', ''))

            if ' — Яндекс' in Window(convertstring).title:

                assert original_browser_name != Window(convertstring).title

    except AssertionError:

        write_bugreport('3','3','сайт загружается','не удалось загрузить сайт')

        assert False, 'Шаг 3 - ПРОВАЛ, Причина: не удалось загрузить сайт'

# Шаг 4: нажать на кнопку "Установить приложение" в адресной строке браузера
def test_step4():
    """Шаг 4: Ожидаемый результат - появляется баллун с выбором режима установки"""

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(2)

    mouse.move(1808, 56)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(2)

# Шаг 5: выбрать пункт "Всплывающее окно"
def test_step5():
    """Шаг 5: Ожидаемый результат - иконка приложения вылетела из адресной строки и прилетела в боковую панель"""

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(2)

    mouse.move(1757, 162)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(20)

# Шаг 6: кликнуть по иконе Вконтакте в боковой панели
def test_step6():
    """Шаг 6: Ожидаемый результат - приложение запускается в виде всплывающего окна"""

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(2)

    mouse.move(25, 840)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(20)

# Шаг 7: нажать на кнопку со стрелкой в шапке всплывающего окна и переключить отображение в режим "Отдельное окно"
def test_step7():
    """Шаг 7: Ожидаемый результат - приложение переключается в режим отдельного окна"""

    mouse = Controller()

    mouse.position = (0,0)

    time.sleep(2)

    mouse.move(423, 104)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    time.sleep(6)

    mouse.position = (0,0)

    time.sleep(2)

    mouse.move(406, 171)

    mouse.press(Button.left)
    time.sleep(0.36)
    mouse.release(Button.left)

    # перебираем словарь и конвертируем id окон в integer для параметра Window и сравниваем кол-во окон
    try:
        window_title = ''

        for i in Window.get_matching_windows(executable='browser.exe'):

            convertstring = str.replace(str(i), 'Win32Window(handle=', '')

            convertstring = int(str.replace(convertstring, ')', ''))

            if 'ВКонтакте' in Window(convertstring).title:

                window_title = Window(convertstring).title

        assert len(Window.get_matching_windows(executable='browser.exe', title=window_title)) != 0

    except AssertionError:

        write_bugreport('3','7','приложение переключается в режим отдельного окна','окно не запустилось')

        assert False, 'Шаг 7 - ПРОВАЛ, Причина: окно не запустилось'